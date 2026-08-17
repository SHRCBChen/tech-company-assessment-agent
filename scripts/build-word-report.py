from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


CN_BODY = "宋体"
CN_HEADING = "黑体"
LATIN_FONT = "Times New Roman"
BODY_SIZE = 12


def set_style_font(style, east_asia: str, latin: str, size: float, bold: bool = False):
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), latin)


def set_run_font(run, east_asia: str = CN_BODY, latin: str = LATIN_FONT, size: float = BODY_SIZE, bold=None):
    run.font.name = latin
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), latin)


def get_or_add_style(doc: Document, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, style_type)


def set_keep_with_next(paragraph, value=True):
    ppr = paragraph._p.get_or_add_pPr()
    node = ppr.find(qn("w:keepNext"))
    if value and node is None:
        ppr.append(OxmlElement("w:keepNext"))
    elif not value and node is not None:
        ppr.remove(node)


def set_widow_control(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    if ppr.find(qn("w:widowControl")) is None:
        ppr.append(OxmlElement("w:widowControl"))


def next_id(elements, tag: str, attr: str) -> int:
    ids = []
    for el in elements.findall(qn(tag)):
        value = el.get(qn(attr))
        if value and value.isdigit():
            ids.append(int(value))
    return max(ids, default=0) + 1


def add_numbering_definition(doc: Document, level_text: str, left: int, hanging: int) -> int:
    numbering = doc.part.numbering_part.element
    abstract_id = next_id(numbering, "w:abstractNum", "w:abstractNumId")
    num_id = next_id(numbering, "w:num", "w:numId")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "chineseCounting")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), level_text)
    lvl.append(lvl_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "space")
    lvl.append(suffix)
    ppr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    ppr.append(ind)
    lvl.append(ppr)
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), LATIN_FONT)
    rfonts.set(qn("w:eastAsia"), CN_HEADING if level_text.endswith("、") else CN_BODY)
    rpr.append(rfonts)
    lvl.append(rpr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_report_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_id = next_id(numbering, "w:abstractNum", "w:abstractNumId")
    num_id = next_id(numbering, "w:num", "w:numId")
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)

    for ilvl, level_text, left, hanging, font_name in [
        (0, "%1、", 0, 0, CN_HEADING),
        (1, "（%2）", 960, 480, CN_BODY),
    ]:
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(ilvl))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "chineseCounting")
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), level_text)
        lvl.append(lvl_text)
        if ilvl == 1:
            restart = OxmlElement("w:lvlRestart")
            restart.set(qn("w:val"), "0")
            lvl.append(restart)
        suffix = OxmlElement("w:suff")
        suffix.set(qn("w:val"), "space")
        lvl.append(suffix)
        ppr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        ppr.append(ind)
        lvl.append(ppr)
        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), LATIN_FONT)
        rfonts.set(qn("w:hAnsi"), LATIN_FONT)
        rfonts.set(qn("w:eastAsia"), font_name)
        rpr.append(rfonts)
        lvl.append(rpr)
        abstract.append(lvl)

    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int, ilvl: int = 0):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl_value = ilvl
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(ilvl_value))
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.append(ilvl)
    numpr.append(numid)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_run_font(run, size=9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(15)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(8)

    normal = doc.styles["Normal"]
    set_style_font(normal, CN_BODY, LATIN_FONT, BODY_SIZE)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5
    pf.first_line_indent = Pt(BODY_SIZE * 2)

    title = get_or_add_style(doc, "报告标题")
    set_style_font(title, CN_BODY, LATIN_FONT, 22, bold=True)
    pf = title.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(0)
    pf.space_after = Pt(12)
    pf.line_spacing = 1.0
    pf.first_line_indent = Pt(0)

    rating = get_or_add_style(doc, "评级摘要")
    set_style_font(rating, CN_BODY, LATIN_FONT, BODY_SIZE, bold=True)
    pf = rating.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(0)
    pf.space_after = Pt(9)
    pf.line_spacing = 1.5
    pf.first_line_indent = Pt(0)

    note = get_or_add_style(doc, "报告说明")
    set_style_font(note, CN_BODY, LATIN_FONT, 10.5)
    pf = note.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(9)
    pf.line_spacing = 1.0
    pf.first_line_indent = Pt(0)
    pf.left_indent = Pt(12)
    pf.right_indent = Pt(12)

    heading = get_or_add_style(doc, "一级标题")
    set_style_font(heading, CN_HEADING, LATIN_FONT, BODY_SIZE, bold=False)
    pf = heading.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(9)
    pf.space_after = Pt(9)
    pf.line_spacing = 1.5
    pf.first_line_indent = Pt(0)
    pf.keep_with_next = True

    subitem = get_or_add_style(doc, "二级编号正文")
    set_style_font(subitem, CN_BODY, LATIN_FONT, BODY_SIZE)
    pf = subitem.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5
    pf.first_line_indent = Pt(0)

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    add_page_number(p)


def clean_inline_markdown(text: str) -> str:
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text).strip()


def build_docx(source: Path, output: Path):
    lines = source.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)
    report_num = add_report_numbering(doc)

    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="报告标题")
            run = p.add_run(clean_inline_markdown(line[2:]))
            set_run_font(run, CN_BODY, LATIN_FONT, 22, bold=True)
            set_keep_with_next(p)
            continue
        if line.startswith("## "):
            title = re.sub(r"^[一二三四五六七八九十]+、\s*", "", clean_inline_markdown(line[3:]))
            p = doc.add_paragraph(style="一级标题")
            apply_numbering(p, report_num, 0)
            run = p.add_run(title)
            set_run_font(run, CN_HEADING, LATIN_FONT, BODY_SIZE, bold=False)
            set_keep_with_next(p)
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="二级编号正文")
            apply_numbering(p, report_num, 1)
            content = clean_inline_markdown(line[2:])
            if "：" in content:
                label, detail = content.split("：", 1)
                r1 = p.add_run(f"{label}：")
                set_run_font(r1, CN_BODY, LATIN_FONT, BODY_SIZE, bold=True)
                r2 = p.add_run(detail)
                set_run_font(r2, CN_BODY, LATIN_FONT, BODY_SIZE)
            else:
                run = p.add_run(content)
                set_run_font(run)
            set_widow_control(p)
            continue
        if line.startswith("> "):
            p = doc.add_paragraph(style="报告说明")
            run = p.add_run(clean_inline_markdown(line[2:]))
            set_run_font(run, CN_BODY, LATIN_FONT, 10.5)
            set_widow_control(p)
            continue
        content = clean_inline_markdown(line)
        if raw.strip().startswith("**") and raw.strip().endswith("**"):
            p = doc.add_paragraph(style="评级摘要")
            run = p.add_run(content)
            set_run_font(run, CN_BODY, LATIN_FONT, BODY_SIZE, bold=True)
            set_keep_with_next(p)
        else:
            p = doc.add_paragraph(style="Normal")
            run = p.add_run(content)
            set_run_font(run)
            set_widow_control(p)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = source.stem.replace("｜", " - ")
    doc.core_properties.subject = "科创企业初评报告"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.save(output)


def main():
    parser = argparse.ArgumentParser(description="将企业初评Markdown转换为统一格式Word报告")
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build_docx(args.source.resolve(), args.output.resolve())
    print(args.output.resolve())


if __name__ == "__main__":
    main()
